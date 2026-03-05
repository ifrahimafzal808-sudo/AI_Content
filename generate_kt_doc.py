"""
Generate a professional Word (.docx) Knowledge Transfer Document
for the AI Content Copilot platform.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Helpers

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, header_color="1F2937", alt_row_color="F3F4F6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, alt_row_color)
    return table

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_note_box(doc, text, prefix="Note: "):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_prefix = p.add_run(prefix)
    run_prefix.bold = True
    run_prefix.italic = True
    run_prefix.font.size = Pt(10)
    run_prefix.font.color.rgb = RGBColor(180, 83, 9)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 53, 15)
    return p


# Build the Document

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

# ===================================================================
#  COVER PAGE
# ===================================================================

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI Content Copilot")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(79, 70, 229)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Knowledge Transfer Document")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(107, 114, 128)

doc.add_paragraph()

divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run("_" * 50)
run.font.color.rgb = RGBColor(209, 213, 219)

doc.add_paragraph()

# Simple one-liner instead of formal metadata block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Everything you need to know about the platform, how it's hosted, and how to keep it running.")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(107, 114, 128)

doc.add_page_break()

# ===================================================================
#  COVER NOTE
# ===================================================================

add_heading(doc, "Cover Note", level=1)

add_para(doc, (
    "I've put together detailed setup docs that explain how the frontend is structured, "
    "how to run it locally, and how to make changes or add new features."
), size=11)

add_para(doc, (
    "The frontend's built using standard React, so any React developer can jump in and "
    "maintain it. With proper documentation and a decent AI coding assistant (like GPT-5 "
    "or Claude), making updates becomes a lot easier than it sounds."
), size=11)

add_para(doc, (
    "If you're specifically looking for a no-code approach for future tweaks, you could "
    "check out something like Lovable (subscription-based, lets you generate frontend code "
    "through prompts). But honestly, for the kind of AI integrations and structured workflows "
    "we've got here, a coded frontend is the more practical long-term choice."
), size=11)

doc.add_page_break()

# ===================================================================
#  TABLE OF CONTENTS
# ===================================================================

add_heading(doc, "Table of Contents", level=1)

toc_items = [
    "1.  What Is This System?",
    "2.  How It Works (Big Picture)",
    "3.  Where Everything Lives",
    "4.  Visiting the Live Website",
    "5.  The Website Pages",
    "6.  How the Backend Works (n8n)",
    "7.  Environment Variables (Settings)",
    "8.  Making Updates to the Website",
    "9.  How Vercel Deployments Work",
    "10. The Code Repository (GitHub)",
    "11. Common Tasks",
    "12. Troubleshooting",
    "13. Glossary",
    "14. Contact and Credentials",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_page_break()

# ===================================================================
#  SECTION 1: What Is This System?
# ===================================================================

add_heading(doc, "1. What Is This System?", level=1)

add_para(doc, (
    "The AI Content Copilot is an internal web-based platform that helps your team "
    "create written content using AI. Think of it as a content production line."
), size=11)

add_para(doc, "Here's how the process works:", size=11)

steps = [
    "You type in a keyword (a topic you want to write about)",
    "AI analyses the keyword and scores it (called a PAVE score) to tell you if it's worth writing about",
    "You approve the keyword and ask AI to create a detailed content plan (an \"Authority Brief\")",
    "AI writes the full article as a professional HTML web page with a cover image",
    "The article gets saved for publishing",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{i}. {step}")
    run.font.size = Pt(10)

add_para(doc, (
    "The whole thing is automated. You just guide it by picking keywords and approving content."
), size=11, space_after=12)

# ===================================================================
#  SECTION 2: How It Works
# ===================================================================

add_heading(doc, "2. How It Works (Big Picture)", level=1)

add_para(doc, "The system has three main parts that talk to each other:", size=11)

add_styled_table(doc,
    ["Part", "What It Does", "Where It Lives"],
    [
        ["Website (Frontend)", "The visual interface you click and interact with", "Hosted on Vercel"],
        ["Brain (Backend)", "The automation engine that talks to AI, searches the web, and stores data", "Hosted on n8n Cloud"],
        ["Data Storage", "Where all your content, brands, and scores are saved", "SharePoint (your existing Microsoft setup)"],
    ]
)

add_para(doc, "", size=6)
add_para(doc, "The flow looks like this:", size=11)
add_para(doc, "You (User)  >  Website (Vercel)  >  Backend (n8n Cloud)  >  SharePoint + AI Services", size=11, color=(79, 70, 229))

add_para(doc, (
    "When you click a button on the website, it sends a request to the n8n backend, "
    "which does the heavy lifting (calling AI, searching the web, reading/writing data "
    "in SharePoint) and sends the results back to the website."
), size=11, space_after=12)

# ===================================================================
#  SECTION 3: Where Everything Lives
# ===================================================================

add_heading(doc, "3. Where Everything Lives", level=1)

add_heading(doc, "3.1 Website (Vercel)", level=2)
add_bullet(doc, "Vercel is a cloud service that hosts websites. Your site is already live there.")
add_bullet(doc, "Live URL: https://ai-content-swart.vercel.app")
add_bullet(doc, "To access the Vercel dashboard, go to vercel.com and log in with the credentials from LastPass.")

add_heading(doc, "3.2 Backend (n8n Cloud)", level=2)
add_bullet(doc, "n8n is a visual automation tool where all the behind-the-scenes logic runs.")
add_bullet(doc, "n8n Cloud URL: https://vitcornu.app.n8n.cloud")
add_bullet(doc, "Log in using the credentials from LastPass.")

add_heading(doc, "3.3 Code (GitHub)", level=2)
add_bullet(doc, "GitHub is where all the source code for the website is stored.")
add_bullet(doc, "Repository: https://github.com/ifrahimafzal808-sudo/AI_Content.git")

add_heading(doc, "3.4 Credentials (LastPass)", level=2)
add_para(doc, "All passwords and login details are stored in LastPass.", size=11)
add_para(doc, "To access LastPass, use the credentials in the shared Google Doc:", size=11)
add_para(doc, "https://docs.google.com/document/d/10ft7sjxjXSmz1JJPT7vM7cSifzWQdbvrUGCYbDB8EiE/edit?usp=sharing", size=10, color=(79, 70, 229))

# ===================================================================
#  SECTION 4: Visiting the Live Website
# ===================================================================

add_heading(doc, "4. Visiting the Live Website", level=1)

add_heading(doc, "4.1 Opening the Website", level=2)
add_para(doc, "Just open your browser and go to:", size=11)
add_para(doc, "https://ai-content-swart.vercel.app", bold=True, size=12, color=(79, 70, 229))
add_para(doc, (
    "You'll see a landing page with the title \"AI-Powered Content Strategy Platform\" "
    "and a button that says \"Go to Dashboard\"."
), size=11)

add_heading(doc, "4.2 Managing the Website on Vercel", level=2)
add_para(doc, "If you need to check the status, view error logs, or change settings:", size=11)

manage_steps = [
    "Go to https://vercel.com",
    "Log in using the credentials from LastPass",
    "You'll see your project listed, click on it",
    "From here you can see if the website is up, view recent deployments, check error logs, and manage environment variables"
]
for i, step in enumerate(manage_steps, 1):
    add_bullet(doc, f"{i}. {step}")

# ===================================================================
#  SECTION 5: The Website Pages
# ===================================================================

add_heading(doc, "5. The Website Pages", level=1)
add_para(doc, "The website has 7 pages. Here's what each one does:", size=11)

add_heading(doc, "5.1 Home Page ( / )", level=2)
add_para(doc, (
    "A landing page with the title \"AI-Powered Content Strategy Platform\" "
    "and three feature cards explaining PAVE Scoring, Authority Briefs, and Content Production."
), size=11)
add_para(doc, "Click \"Go to Dashboard\" to get into the main app.", size=11)

add_heading(doc, "5.2 Dashboard ( /dashboard )", level=2)
add_para(doc, "An overview of your content production with stats:", size=11)
for stat in ["Ideas Analyzed", "Approved for Briefing", "In Production", "Published"]:
    add_bullet(doc, stat)
add_para(doc, "Below the stats there's a Recent Activity feed showing the latest actions.", size=11)

add_heading(doc, "5.3 Analyze Keyword ( /analyze )", level=2)
add_para(doc, "This is where you evaluate new content ideas. There are two inputs:", size=11)
add_bullet(doc, "Select Brand - a dropdown to choose which brand you're writing for")
add_bullet(doc, "Keyword - a text box where you type the topic")

add_para(doc, "How to use it:", size=11)
analyze_steps = [
    "Select a brand from the dropdown",
    "Type your keyword",
    "Click \"Analyze with AI\"",
    "Wait while the AI searches the web and evaluates your keyword",
    "You'll see PAVE Scores (four scores, 1 to 5)",
    "You can adjust the scores if you disagree with the AI",
    "Click \"Save to Backlog\" to save it"
]
for i, step in enumerate(analyze_steps, 1):
    add_bullet(doc, f"{i}. {step}")

add_para(doc, "The PAVE scores explained:", size=11)
add_styled_table(doc,
    ["Score", "Full Name", "What It Measures"],
    [
        ["P", "Profitability", "How likely this topic is to generate revenue"],
        ["A", "Authority", "How well your brand can be authoritative on this topic"],
        ["V", "Volume", "How many people search for this topic (always 0 for now, requires manual check)"],
        ["E", "Effort", "How difficult it'll be to compete on this topic"],
    ]
)

add_heading(doc, "5.4 Content Backlog ( /backlog )", level=2)
add_para(doc, (
    "A table listing all your saved keywords and their current status. "
    "You can filter by clicking status buttons at the top."
), size=11)
add_para(doc, "Available filters:", size=11)
for status in ["All", "PAVE Scored", "Approved for Briefing", "Briefed", "In Production", "Published"]:
    add_bullet(doc, status)

add_heading(doc, "5.5 Create Brief ( /brief )", level=2)
add_para(doc, "A step-by-step wizard to create an Authority Brief (a detailed content plan):", size=11)
brief_steps = [
    "Select a keyword from the dropdown (only shows items marked \"Approved for Briefing\")",
    "Click \"Get Format Recommendation\" - the AI suggests what type of article to write",
    "Optionally add a Differentiator (what makes your brand special) and SME Input (expert notes)",
    "Click \"Generate Authority Brief\" - the AI creates a detailed content plan",
    "Review it and click \"Approve and Save Brief\" to save",
]
for i, step in enumerate(brief_steps, 1):
    add_bullet(doc, f"{i}. {step}")

add_heading(doc, "5.6 Produce Content ( /produce )", level=2)
add_para(doc, "This page generates the full HTML article:", size=11)
produce_steps = [
    "Select a briefed item from the dropdown",
    "Click \"Start Content Production\"",
    "A progress indicator shows 5 steps: Preparing Context > Generating HTML Article > Generating Cover Image > Assembling Final HTML > Saving to SharePoint",
    "When it's done, you'll see a success screen with a link to view the article"
]
for i, step in enumerate(produce_steps, 1):
    add_bullet(doc, f"{i}. {step}")

add_note_box(doc, (
    "The \"Saving to SharePoint\" step (step 5) isn't fully implemented "
    "in the backend yet. That's the remaining dev work."
), prefix="Heads up: ")

add_heading(doc, "5.7 Published Gallery ( /published )", level=2)
add_para(doc, (
    "A gallery grid showing all published articles with the title, "
    "brand name, published date, and PAVE total score."
), size=11)

# ===================================================================
#  SECTION 6: How the Backend Works
# ===================================================================

add_heading(doc, "6. How the Backend Works (n8n)", level=1)

add_para(doc, (
    "The n8n backend has automated workflows (think of them like recipes) that "
    "the website calls when you click buttons. Here are the main ones:"
), size=11)

add_styled_table(doc,
    ["Workflow Name", "What It Does", "Triggered When"],
    [
        ["GetActiveBrands", "Fetches the list of brands from SharePoint", "You open the Analyze Keyword page"],
        ["GetBrandDetails", "Gets full details for a specific brand", "Creating a brief"],
        ["AnalyzeKeyword_Backend", "Searches the web + asks AI for PAVE scores", "You click \"Analyze with AI\""],
        ["SaveScoredIdea", "Saves the keyword + scores to SharePoint", "You click \"Save to Backlog\""],
        ["FetchAllItems", "Gets all items from the content pipeline", "Opening the Content Backlog page"],
        ["FetchBriefedItem", "Gets items filtered by a specific status", "Filtering by status in backlog"],
        ["FetchItemDetails", "Gets full details for one specific item", "Selecting an item for briefing"],
        ["BriefRecommendation", "AI recommends a content format", "You click \"Get Format Recommendation\""],
        ["GenerateAuthorityBrief", "AI generates a full content brief", "You click \"Generate Authority Brief\""],
        ["UpdateBriefAndStatus", "Saves the approved brief to SharePoint", "You click \"Approve and Save Brief\""],
        ["ProduceContent", "Generates the full HTML article", "You click \"Start Content Production\""],
        ["UpdateContentStatus", "Changes the status of a content item", "Publishing content"],
    ]
)

add_para(doc, "", size=6)
add_heading(doc, "How to Access n8n", level=2)
access_steps = [
    "Go to https://vitcornu.app.n8n.cloud",
    "Log in with the credentials from LastPass",
    "You'll see a list of all workflows",
    "Click any workflow to see its visual flowchart",
]
for i, step in enumerate(access_steps, 1):
    add_bullet(doc, f"{i}. {step}")

add_note_box(doc, "Don't modify workflows unless you've got a developer making changes.", prefix="Warning: ")

# ===================================================================
#  SECTION 7: Environment Variables
# ===================================================================

add_heading(doc, "7. Environment Variables (Settings)", level=1)

add_para(doc, (
    "Environment variables are settings that tell the website where to connect. "
    "Think of them like an address book."
), size=11)

add_heading(doc, "Current Settings", level=2)

add_styled_table(doc,
    ["Setting Name", "What It Controls", "Current Value"],
    [
        ["VITE_N8N_API_URL", "Where the website sends requests to the backend", "/api/n8n (Vercel routes this to n8n Cloud)"],
        ["VITE_N8N_API_KEY", "A secret key to authenticate with n8n", "Stored in the .env file"],
        ["VITE_USE_MOCK_DATA", "If set to 'true', shows fake test data", "Currently 'false' (using real data)"],
    ]
)

add_para(doc, "", size=6)

add_heading(doc, "Where These Settings Live", level=2)
add_bullet(doc, "In the code: a file called .env inside the frontend folder")
add_bullet(doc, "On Vercel: in the dashboard under Settings > Environment Variables")

add_note_box(doc, (
    "The settings in Vercel must match the .env file. "
    "If you change a setting, update it in both places."
), prefix="Important: ")

add_heading(doc, "How to Change Settings on Vercel", level=2)
vercel_env_steps = [
    "Go to https://vercel.com and log in",
    "Click on your project",
    "Go to Settings (tab at the top)",
    "Click Environment Variables in the left sidebar",
    "You'll see a list of all settings. You can edit, add, or remove them",
    "After making changes, you'll need to redeploy for the changes to take effect"
]
for i, step in enumerate(vercel_env_steps, 1):
    add_bullet(doc, f"{i}. {step}")

# ===================================================================
#  SECTION 8: Making Updates
# ===================================================================

add_heading(doc, "8. Making Updates to the Website", level=1)

add_para(doc, "If you or a developer needs to make changes:", size=11)

add_heading(doc, "Step 1: Edit the Code", level=2)
add_para(doc, "The code lives in the GitHub repository. A developer would:", size=11)
edit_steps = [
    "Download (clone) the code from GitHub",
    "Make the necessary changes on their computer",
    "Test the changes locally by running the dev server",
    "Push the changes back to GitHub"
]
for i, step in enumerate(edit_steps, 1):
    add_bullet(doc, f"{i}. {step}")

add_heading(doc, "Step 2: Automatic Deployment", level=2)
add_para(doc, (
    "Once changes are pushed to GitHub, Vercel automatically picks up the update "
    "and deploys a new version. This usually takes 1-2 minutes."
), size=11)

add_heading(doc, "Running the Website Locally", level=2)
add_para(doc, "If a developer needs to run the website on their own machine for testing:", size=11)

local_steps = [
    "Download the code: git clone https://github.com/ifrahimafzal808-sudo/AI_Content.git",
    "Go to the frontend folder: cd frontend",
    "Install dependencies: npm install",
    "Start the dev server: npm run dev",
    "Open the website at the URL shown in terminal (usually http://localhost:5173)",
]
for i, step in enumerate(local_steps, 1):
    add_bullet(doc, f"{i}. {step}")

add_note_box(doc, (
    "For testing without the n8n backend, set VITE_USE_MOCK_DATA to 'true' "
    "in the .env file. This shows fake data so you can test the interface."
), prefix="Tip: ")

# ===================================================================
#  SECTION 9: Vercel Deployment
# ===================================================================

add_heading(doc, "9. How Vercel Deployments Work", level=1)

add_para(doc, "Vercel's connected to the GitHub repository. Here's what happens when code is updated:", size=11)

deploy_flow = [
    "Developer pushes code to GitHub",
    "Vercel detects the change automatically",
    "Vercel builds the new version (compiles the code)",
    "Vercel deploys it to the live URL",
    "The website is updated (usually within 1-2 minutes)"
]
for i, step in enumerate(deploy_flow, 1):
    p = doc.add_paragraph()
    run_n = p.add_run(f"  {i}.  ")
    run_n.font.size = Pt(10)
    run_n.font.color.rgb = RGBColor(79, 70, 229)
    run_t = p.add_run(step)
    run_t.font.size = Pt(10)
    if i < len(deploy_flow):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_arrow = p2.add_run("v")
        run_arrow.font.size = Pt(12)
        run_arrow.font.color.rgb = RGBColor(156, 163, 175)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)

add_heading(doc, "How to See Deployment History", level=2)
deploy_history_steps = [
    "Go to your project in the Vercel dashboard",
    "Click the Deployments tab",
    "You'll see a list of all deployments with dates, status (Ready, Error, Building), and which code change triggered it",
    "Click any deployment to see details and logs"
]
for i, step in enumerate(deploy_history_steps, 1):
    add_bullet(doc, f"{i}. {step}")

add_heading(doc, "How to Manually Redeploy", level=2)
redeploy_steps = [
    "Go to your project in the Vercel dashboard",
    "Click the Deployments tab",
    "Find the most recent successful deployment",
    "Click the three-dot menu next to it",
    "Select Redeploy",
    "Confirm the action"
]
for i, step in enumerate(redeploy_steps, 1):
    add_bullet(doc, f"{i}. {step}")

# ===================================================================
#  SECTION 10: Code Repository
# ===================================================================

add_heading(doc, "10. The Code Repository (GitHub)", level=1)

add_heading(doc, "Repository URL", level=2)
add_para(doc, "https://github.com/ifrahimafzal808-sudo/AI_Content.git", bold=True, size=11, color=(79, 70, 229))

add_heading(doc, "Important Files and Folders", level=2)

add_styled_table(doc,
    ["File / Folder", "What It Is"],
    [
        [".env", "Settings file (n8n URL, API key)"],
        ["vercel.json", "Vercel routing config - routes /api/n8n to n8n Cloud"],
        ["package.json", "List of software dependencies"],
        ["vite.config.js", "Build tool configuration"],
        ["index.html", "The main HTML shell"],
        ["src/App.jsx", "Defines all the pages and their web addresses"],
        ["src/main.jsx", "The application entry point"],
        ["src/index.css", "All visual styling (colours, fonts, spacing)"],
        ["src/components/", "Reusable visual building blocks"],
        ["src/hooks/", "Helper code used across pages"],
        ["src/pages/", "The 7 individual pages (see Section 5)"],
        ["src/services/api.js", "Code that talks to the n8n backend"],
    ]
)

add_para(doc, "", size=6)
add_heading(doc, "The Vercel Routing File (vercel.json)", level=2)
add_para(doc, (
    "This file tells Vercel how to route requests. It's got one important rule: "
    "any request to /api/n8n/... gets forwarded to https://vitcornu.app.n8n.cloud/... "
    "So the website never talks directly to n8n. Vercel acts as a middleman, which is more secure."
), size=11)

# ===================================================================
#  SECTION 11: Common Tasks
# ===================================================================

add_heading(doc, "11. Common Tasks", level=1)

tasks = [
    ("\"The website is down\"",
     [
         "Go to Vercel dashboard, check if the last deployment failed",
         "If it shows an error, click on it to see the error logs",
         "Try redeploying the last successful version (see Section 9)",
         "If that doesn't help, check if n8n Cloud is running at https://vitcornu.app.n8n.cloud"
     ]),
    ("\"Data isn't showing on the website\"",
     [
         "Check if VITE_USE_MOCK_DATA is set to 'false' (should be for real data)",
         "Go to n8n Cloud and check if the workflows are active (turned on)",
         "Check if SharePoint is accessible and has data"
     ]),
    ("\"I need to change the n8n backend URL\"",
     [
         "Update the VITE_N8N_API_URL in the .env file",
         "Also update it in Vercel dashboard > Settings > Environment Variables",
         "Update the vercel.json file to point the proxy to the new URL",
         "Redeploy on Vercel"
     ]),
    ("\"I want to test without the live backend\"",
     [
         "Change VITE_USE_MOCK_DATA to 'true' in the .env file",
         "The website will show sample/fake data so you can explore the interface"
     ]),
]

for task_title, task_steps in tasks:
    add_heading(doc, task_title, level=2)
    for i, step in enumerate(task_steps, 1):
        add_bullet(doc, f"{i}. {step}")

# ===================================================================
#  SECTION 12: Troubleshooting
# ===================================================================

add_heading(doc, "12. Troubleshooting", level=1)

add_styled_table(doc,
    ["Problem", "Likely Cause", "What to Do"],
    [
        ["Website shows a blank page", "JavaScript error or build failure", "Check Vercel deployment logs for errors"],
        ["\"Analysis failed\" when analyzing", "n8n workflow isn't running, or API key is wrong", "Check n8n Cloud, make sure AnalyzeKeyword_Backend workflow is active"],
        ["No brands in the dropdown", "n8n can't reach SharePoint, or GetActiveBrands workflow is off", "Check n8n Cloud workflows are active and SharePoint is accessible"],
        ["Backlog table is empty", "No data in SharePoint, or mock mode is on", "Make sure VITE_USE_MOCK_DATA is 'false' and SharePoint has data"],
        ["\"API request failed\" error", "Network issue or wrong API URL", "Check VITE_N8N_API_URL is correct and n8n Cloud is online"],
        ["Buttons don't work", "JavaScript error", "Open browser developer tools (F12) > Console tab > check for red errors"],
        ["GitHub changes not showing up", "Vercel auto-deploy might've failed", "Go to Vercel > Deployments tab to check build status"],
    ]
)

# ===================================================================
#  SECTION 13: Glossary
# ===================================================================

add_heading(doc, "13. Glossary", level=1)

add_styled_table(doc,
    ["Term", "What It Means"],
    [
        ["Frontend", "The visual part of the website, what you see and click on"],
        ["Backend", "The behind-the-scenes logic that processes data and calls AI"],
        ["Vercel", "The cloud service that hosts the website"],
        ["n8n", "The automation platform that runs all the backend workflows"],
        ["GitHub", "Where the source code is stored and tracked"],
        ["SharePoint", "Microsoft's cloud storage where brand configs and content data lives"],
        ["Deployment", "Putting a new version of the website online"],
        ["Environment Variables", "Settings that tell the website where to connect"],
        ["PAVE Score", "Profitability, Authority, Volume, Effort - keyword quality scores"],
        ["Authority Brief", "A detailed plan created by AI to guide content production"],
        ["Webhook", "A web address that triggers a specific backend workflow"],
        ["Repository (Repo)", "The online folder on GitHub containing all code files"],
        ["Clone", "Downloading a copy of the code from GitHub to a local machine"],
        ["Mock Data", "Fake sample data used for testing without the real backend"],
        ["API", "A set of rules that lets the website talk to the backend"],
        ["HTML", "The code language used to structure web pages"],
        ["React", "The framework used to build the website's interface"],
        ["Vite", "A tool that helps developers build and test React websites quickly"],
        ["LastPass", "A password management tool where all shared credentials are stored"],
        ["npm", "A tool developers use to install software packages"],
    ]
)

# ===================================================================
#  SECTION 14: Contact and Credentials
# ===================================================================

add_heading(doc, "14. Contact and Credentials", level=1)

add_heading(doc, "Credentials Access", level=2)
add_para(doc, "All credentials are stored in the shared Google Doc:", size=11)
add_para(doc, "https://docs.google.com/document/d/10ft7sjxjXSmz1JJPT7vM7cSifzWQdbvrUGCYbDB8EiE/edit?usp=sharing",
         size=10, color=(79, 70, 229))

add_para(doc, "This doc contains:", size=11)
add_bullet(doc, "Gmail account credentials for Ibrahim's dedicated project email")
add_bullet(doc, "LastPass master credentials (email and master password)")

add_heading(doc, "Inside LastPass You'll Find", level=2)
lp_items = [
    "n8n Cloud login",
    "Vercel login",
    "SharePoint access details",
    "GitHub account details",
    "Any other project-related passwords"
]
for item in lp_items:
    add_bullet(doc, item)

# Footer
doc.add_paragraph()
divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run("_" * 50)
run.font.color.rgb = RGBColor(209, 213, 219)


# ===================================================================
#  SAVE
# ===================================================================

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "AI_Content_Copilot_Knowledge_Transfer.docx"
)
doc.save(output_path)
print(f"\n  Document saved to:\n    {output_path}")
print(f"    File size: {os.path.getsize(output_path) / 1024:.1f} KB")
