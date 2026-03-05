"""
Generate a professional Word (.docx) Frontend Maintenance & Developer Guide
for the AI Content Copilot platform.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Helpers

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, header_color="1F2937", alt_row_color="F3F4F6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, alt_row_color)
    return table

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_note_box(doc, text, prefix="Note: "):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_prefix = p.add_run(prefix)
    run_prefix.bold = True
    run_prefix.italic = True
    run_prefix.font.size = Pt(10)
    run_prefix.font.color.rgb = RGBColor(180, 83, 9)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 53, 15)
    return p

def add_code_block(doc, code_text, label=None):
    if label:
        p_label = doc.add_paragraph()
        run_l = p_label.add_run(label)
        run_l.font.size = Pt(9)
        run_l.font.color.rgb = RGBColor(107, 114, 128)
        p_label.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
    pPr.append(shading)
    return p


# Build Document

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

# ===================================================================
#  COVER PAGE
# ===================================================================

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI Content Copilot")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(79, 70, 229)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Frontend Maintenance &\nDeveloper Guide")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(107, 114, 128)

doc.add_paragraph()

divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run("_" * 50)
run.font.color.rgb = RGBColor(209, 213, 219)

doc.add_paragraph()

# Simple one-liner instead of formal metadata
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("How the frontend was built, how to maintain it, and how to add or change things.")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(107, 114, 128)

doc.add_page_break()

# ===================================================================
#  TABLE OF CONTENTS
# ===================================================================

add_heading(doc, "Table of Contents", level=1)

toc_items = [
    "1.  What This Guide Covers",
    "2.  Technology Stack",
    "3.  Project Structure (Complete File Map)",
    "4.  How the Pieces Fit Together",
    "5.  The Styling System",
    "6.  How Pages Work",
    "7.  Adding a New Page",
    "8.  Modifying an Existing Page",
    "9.  How the Website Talks to the Backend (API Layer)",
    "10. Adding a New Backend Connection",
    "11. The Navigation System",
    "12. Adding a New Navigation Link",
    "13. Reusable Components",
    "14. Environment Settings",
    "15. Running the Website Locally",
    "16. How Deployments Work",
    "17. Maintaining the App with AI",
    "18. Quick Reference Cheat Sheet",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_page_break()

# ===================================================================
#  SECTION 1: Introduction
# ===================================================================

add_heading(doc, "1. What This Guide Covers", level=1)

add_para(doc, (
    "This doc explains how the AI Content Copilot frontend was built, how all the "
    "pieces fit together, and most importantly, how to add new features, change "
    "existing ones, or fix issues."
), size=11)

add_para(doc, (
    "It's written so that even if you're not a developer, you can understand what "
    "a developer would need to do. You can also hand this to an AI coding assistant "
    "(like GPT-5 or Claude) and it'll have enough context to make changes."
), size=11)

add_para(doc, "What you'll learn:", size=11)
learn_items = [
    "What each file in the project does",
    "How to add a brand-new page to the website",
    "How to change text, labels, or behaviour on existing pages",
    "How to connect a new backend workflow (n8n) to the website",
    "How to add a new link in the sidebar",
    "How to change the look and feel (colours, fonts, spacing)",
    "How to deploy changes",
    "How to use AI coding assistants to maintain the app",
]
for item in learn_items:
    add_bullet(doc, item)

# ===================================================================
#  SECTION 2: Technology Stack
# ===================================================================

add_heading(doc, "2. Technology Stack", level=1)

add_para(doc, (
    "We're using standard, widely-supported technologies here. Any competent "
    "web developer can work on this, and AI coding tools are highly trained "
    "on this exact stack."
), size=11)

add_styled_table(doc,
    ["Technology", "What It Is", "Why We Use It"],
    [
        ["React 19", "Framework for building interactive websites", "The most popular web framework out there. Huge community, easy to hire for"],
        ["Vite 7", "Build tool that compiles the code", "Very fast dev experience, industry standard"],
        ["React Router 7", "Handles page navigation (URL to page)", "Standard way to manage multiple pages in React"],
        ["Vanilla CSS", "Custom styling (colours, layout, fonts)", "Maximum flexibility, no dependency on CSS frameworks"],
        ["Vercel", "Website hosting and deployment", "Free tier available, auto-deploys from GitHub"],
    ]
)

add_para(doc, "", size=6)
add_note_box(doc, (
    "We went with vanilla CSS instead of something like Tailwind because it's simpler "
    "for this project's size and doesn't add another dependency. If the app grows "
    "significantly, a component library might make sense down the line."
), prefix="Why this setup: ")

# ===================================================================
#  SECTION 3: Project Structure
# ===================================================================

add_heading(doc, "3. Project Structure (Complete File Map)", level=1)

add_para(doc, (
    "Here's every file and folder in the project, with a plain description "
    "of what each one does."
), size=11)

add_heading(doc, "Root Files (Configuration)", level=2)
add_styled_table(doc,
    ["File", "What It Does", "When You'd Touch It"],
    [
        [".env", "Stores settings (n8n URL, mock mode)", "When changing which backend to connect to, or toggling mock data"],
        [".env.example", "Template showing what settings are available", "Never, it's just a reference for new devs"],
        ["vercel.json", "Tells Vercel how to route requests (proxies /api/n8n to n8n Cloud)", "When changing the n8n Cloud URL"],
        ["vite.config.js", "Configures the build tool (also sets up the local proxy)", "When changing the n8n Cloud URL for local dev"],
        ["package.json", "Lists all dependencies and scripts (npm install, npm run dev)", "When adding a new library"],
        ["index.html", "The HTML shell that React fills in", "Rarely, only to change the page title or meta tags"],
    ]
)

add_heading(doc, "Source Code (src/ folder)", level=2)

add_heading(doc, "Core Files", level=3)
add_styled_table(doc,
    ["File", "What It Does", "When You'd Touch It"],
    [
        ["src/main.jsx", "The entry point, loads React and renders the app", "Rarely, only if changing app-wide wrappers"],
        ["src/App.jsx", "Defines ALL pages and their URL paths (routing)", "When adding a new page"],
        ["src/index.css", "ALL visual styling: colours, fonts, buttons, cards, layout", "When changing how anything looks"],
    ]
)

add_heading(doc, "Pages (src/pages/)", level=3)
add_para(doc, "Each file here is one full page of the website:", size=11)
add_styled_table(doc,
    ["File", "URL Path", "What The Page Does"],
    [
        ["Home.jsx", "/", "Landing page with hero section and feature cards"],
        ["Dashboard.jsx", "/dashboard", "Overview with stats and recent activity feed"],
        ["AnalyzeKeyword.jsx", "/analyze", "Form to analyze keywords and get PAVE scores from AI"],
        ["ContentBacklog.jsx", "/backlog", "Table showing all saved keywords with status filters"],
        ["CreateBrief.jsx", "/brief", "Step-by-step wizard to create AI authority briefs"],
        ["ProduceContent.jsx", "/produce", "Generate full HTML articles with progress tracking"],
        ["Published.jsx", "/published", "Gallery grid of all published articles"],
    ]
)

add_heading(doc, "Components (src/components/)", level=3)
add_para(doc, "Reusable building blocks used across multiple pages:", size=11)
add_styled_table(doc,
    ["File", "What It Is", "Used By"],
    [
        ["layout/MainLayout.jsx", "The page shell - combines Navbar + Sidebar + page content", "All pages except Home"],
        ["layout/Navbar.jsx", "Top navigation bar with logo and user info", "MainLayout (appears on every page)"],
        ["layout/Sidebar.jsx", "Left sidebar with navigation links", "MainLayout (appears on every page)"],
        ["common/LoadingSpinner.jsx", "A spinning animation shown while data loads", "AnalyzeKeyword, Backlog, Brief, Produce, Published"],
        ["common/PAVEScoreCard.jsx", "The card showing P/A/V/E scores with editable inputs", "AnalyzeKeyword page"],
        ["common/StatusBadge.jsx", "Coloured badge showing item status (e.g. 'Briefed')", "ContentBacklog page"],
    ]
)

add_heading(doc, "Services (src/services/)", level=3)
add_styled_table(doc,
    ["File", "What It Does", "When You'd Touch It"],
    [
        ["api.js", "ALL backend communication, every n8n webhook call lives here", "When adding a new backend feature or changing how data flows"],
    ]
)

add_heading(doc, "Hooks (src/hooks/)", level=3)
add_styled_table(doc,
    ["File", "What It Does", "When You'd Touch It"],
    [
        ["useApi.js", "Helper that handles loading states and errors for API calls", "Rarely, it's a utility used by pages"],
    ]
)

# ===================================================================
#  SECTION 4: Architecture
# ===================================================================

add_heading(doc, "4. How the Pieces Fit Together", level=1)

add_para(doc, "Here's how the website works from top to bottom:", size=11)

flow_steps = [
    ("User visits the website", "The browser loads index.html, which loads main.jsx"),
    ("main.jsx starts React", "It wraps everything in a Router (for page navigation)"),
    ("App.jsx defines the pages", "It maps URL paths to page components (e.g., /analyze goes to AnalyzeKeyword.jsx)"),
    ("MainLayout wraps internal pages", "It adds the Navbar (top bar) and Sidebar (left menu) around page content"),
    ("Pages use api.js to get data", "When a page needs data, it calls a function in api.js"),
    ("api.js talks to n8n", "It sends requests to the n8n backend through the Vercel proxy"),
    ("Data comes back, page updates", "React automatically re-renders with the new data"),
]
for action, detail in flow_steps:
    p = doc.add_paragraph()
    run1 = p.add_run(action + "  >  ")
    run1.font.size = Pt(10)
    run1.font.color.rgb = RGBColor(79, 70, 229)
    run2 = p.add_run(detail)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

add_para(doc, "", size=6)
add_para(doc, (
    "The key thing to know: each page is independent. You can change AnalyzeKeyword.jsx "
    "without affecting ContentBacklog.jsx. The only shared pieces are the layout (Navbar, "
    "Sidebar) and the API service (api.js)."
), size=11)

add_para(doc, (
    "We didn't use a state management library like Redux because the app's simple enough "
    "that it doesn't need one. Each page manages its own state. If the app grows to the "
    "point where pages need to share a lot of data, that might change."
), size=11)

# ===================================================================
#  SECTION 5: Styling System
# ===================================================================

add_heading(doc, "5. The Styling System", level=1)

add_para(doc, (
    "All visual styling lives in one file: src/index.css. It's about 600+ lines, "
    "controlling every colour, font, spacing, animation, and layout in the app."
), size=11)

add_heading(doc, "How Styles Are Organized", level=2)
add_styled_table(doc,
    ["Section", "What It Controls", "Examples"],
    [
        ["CSS Variables (top)", "Colour palette, font sizes, spacing values", "--color-primary, --font-size-sm, --space-4"],
        ["Base Styles", "Default text, body background, links", "Body font, background colour"],
        ["Layout", "Page structure, sidebar width, navbar height", ".main-layout, .sidebar, .navbar"],
        ["Cards & Containers", "Content boxes used across pages", ".card, .table-container"],
        ["Buttons", "All button styles and variants", ".btn-primary, .btn-secondary, .btn-ghost"],
        ["Forms", "Input fields, dropdowns, textareas", ".input, .select, .label"],
        ["Tables", "Data table styling", "table, th, td"],
        ["Badges", "Status indicators (coloured pills)", ".badge-success, .badge-warning"],
        ["Animations", "Fade-ins, spinners, transitions", ".animate-fade-in, .spinner"],
        ["Page-Specific", "Styles for specific pages", ".landing-hero, .gallery-grid, .pave-grid"],
    ]
)

add_para(doc, "", size=6)
add_heading(doc, "How to Change the Colour Scheme", level=2)
add_para(doc, (
    "All colours are defined as CSS Variables at the top of index.css. To change the "
    "colour scheme, you just update these variables and everything updates automatically."
), size=11)

add_code_block(doc, (
    "/* Find these at the top of src/index.css */\n"
    ":root {\n"
    "    --color-primary: #6366F1;       /* Main purple - buttons, links, accents */\n"
    "    --color-primary-hover: #5558E6; /* Button hover state */\n"
    "    --color-success: #10B981;       /* Green - success messages, badges */\n"
    "    --color-warning: #F59E0B;       /* Amber - warnings */\n"
    "    --color-error: #EF4444;         /* Red - errors */\n"
    "    --bg-primary: #0A0A1A;          /* Main background colour */\n"
    "    --text-primary: #FFFFFF;        /* Main text colour */\n"
    "}"
), label="Colour Variables in index.css:")

add_note_box(doc, (
    "To change the brand colour from purple to blue, you'd just swap "
    "--color-primary from #6366F1 to something like #3B82F6. "
    "Every button, link, and accent in the app updates."
), prefix="Example: ")

# ===================================================================
#  SECTION 6: How Pages Work
# ===================================================================

add_heading(doc, "6. How Pages Work", level=1)

add_para(doc, (
    "Every page follows the same pattern. Once you understand it, "
    "you can read (and modify) any page."
), size=11)

add_para(doc, "The pattern:", size=11)

pattern_steps = [
    ("Import tools:", "The page imports React hooks, API functions, and any components it needs"),
    ("Define state:", "Variables that track what's happening (is data loading? what did the user type?)"),
    ("Define actions:", "Functions that run when the user clicks a button"),
    ("Return layout:", "The HTML-like structure that defines what appears on screen"),
]
for i, (step, detail) in enumerate(pattern_steps, 1):
    p = doc.add_paragraph()
    run_n = p.add_run(f"  {i}. ")
    run_n.font.size = Pt(10)
    run_n.font.color.rgb = RGBColor(79, 70, 229)
    run_b = p.add_run(step + " ")
    run_b.font.size = Pt(10)
    run_d = p.add_run(detail)
    run_d.font.size = Pt(10)

add_para(doc, "", size=4)
add_code_block(doc, (
    "// Simplified page structure (every page looks like this)\n\n"
    "import { useState } from 'react'           // 1. Import tools\n"
    "import { api } from '../services/api'\n\n"
    "export default function MyPage() {\n"
    "    const [data, setData] = useState(null)  // 2. Define state\n\n"
    "    const handleClick = async () => {       // 3. Define actions\n"
    "        const result = await api.someFunction()\n"
    "        setData(result)\n"
    "    }\n\n"
    "    return (                                // 4. Layout\n"
    "        <div>\n"
    "            <h1>My Page Title</h1>\n"
    "            <button onClick={handleClick}>Do Something</button>\n"
    "            {data && <p>{data.message}</p>}\n"
    "        </div>\n"
    "    )\n"
    "}"
), label="Page Pattern Template:")

# ===================================================================
#  SECTION 7: Adding a New Page
# ===================================================================

add_heading(doc, "7. Adding a New Page", level=1)

add_para(doc, (
    "Adding a page requires changes in exactly 3 files. That's it."
), size=11)

add_heading(doc, "Step 1: Create the Page File", level=2)
add_para(doc, "Create a new file in src/pages/. Say you want a \"Settings\" page:", size=11)
add_code_block(doc, (
    "// File: src/pages/Settings.jsx\n\n"
    "export default function Settings() {\n"
    "    return (\n"
    "        <div className=\"animate-fade-in\">\n"
    "            <div className=\"page-header\">\n"
    "                <h1 className=\"page-title\">Settings</h1>\n"
    "                <p className=\"page-subtitle\">Manage your account settings</p>\n"
    "            </div>\n\n"
    "            <div className=\"card\">\n"
    "                <h3>Your Settings Go Here</h3>\n"
    "                <p>Add forms, toggles, or whatever you need.</p>\n"
    "            </div>\n"
    "        </div>\n"
    "    )\n"
    "}"
), label="Create: src/pages/Settings.jsx")

add_heading(doc, "Step 2: Register the Route", level=2)
add_para(doc, "Open src/App.jsx and add your new page:", size=11)
add_code_block(doc, (
    "// In src/App.jsx, add these two things:\n\n"
    "// 1. Add this import at the top:\n"
    "import Settings from './pages/Settings'\n\n"
    "// 2. Add this route inside the <Route element={<MainLayout />}> block:\n"
    "<Route path=\"/settings\" element={<Settings />} />"
), label="Edit: src/App.jsx")

add_heading(doc, "Step 3: Add to the Sidebar", level=2)
add_para(doc, "Open src/components/layout/Sidebar.jsx and add a new link:", size=11)
add_code_block(doc, (
    "// In src/components/layout/Sidebar.jsx\n"
    "// Add this entry to the 'navigation' array:\n\n"
    "const navigation = [\n"
    "    { name: 'Dashboard', href: '/dashboard', icon: 'chart' },\n"
    "    { name: 'Analyze Keyword', href: '/analyze', icon: 'search' },\n"
    "    { name: 'Content Backlog', href: '/backlog', icon: 'list' },\n"
    "    { name: 'Create Brief', href: '/brief', icon: 'edit' },\n"
    "    { name: 'Produce Content', href: '/produce', icon: 'zap' },\n"
    "    { name: 'Published', href: '/published', icon: 'rocket' },\n"
    "    { name: 'Settings', href: '/settings', icon: 'gear' }   // <-- NEW\n"
    "]"
), label="Edit: src/components/layout/Sidebar.jsx")

add_note_box(doc, (
    "That's it. Three files, three simple changes. Push to GitHub and Vercel "
    "auto-deploys the update."
), prefix="Done! ")

# ===================================================================
#  SECTION 8: Modifying an Existing Page
# ===================================================================

add_heading(doc, "8. Modifying an Existing Page", level=1)

add_para(doc, "Common modifications and where to make them:", size=11)

add_styled_table(doc,
    ["What You Want to Change", "File to Edit", "What to Look For"],
    [
        ["Page title or description", "src/pages/[PageName].jsx", "Look for <h1 className=\"page-title\"> and <p className=\"page-subtitle\">"],
        ["Button text", "src/pages/[PageName].jsx", "Look for <button> elements"],
        ["Form fields (add/remove/rename)", "src/pages/[PageName].jsx", "Look for <input>, <select>, or <textarea> elements"],
        ["The status filter options on Backlog", "src/pages/ContentBacklog.jsx", "Edit the STATUS_FILTERS array at the top"],
        ["Dashboard statistics", "src/pages/Dashboard.jsx", "Edit the <div className=\"stat-card\"> sections"],
        ["Production progress steps", "src/pages/ProduceContent.jsx", "Edit the PRODUCTION_STEPS array at the top"],
        ["PAVE scoring display", "src/components/common/PAVEScoreCard.jsx", "Edit the score labels, ranges, or layout"],
        ["Colours, fonts, spacing", "src/index.css", "Edit the CSS variables at the top"],
        ["Sidebar navigation links", "src/components/layout/Sidebar.jsx", "Edit the navigation array"],
        ["Navbar logo or title", "src/components/layout/Navbar.jsx", "Edit the text in the navbar-brand div"],
    ]
)

# ===================================================================
#  SECTION 9: API Layer
# ===================================================================

add_heading(doc, "9. How the Website Talks to the Backend (API Layer)", level=1)

add_para(doc, (
    "All communication between the website and n8n goes through one file: "
    "src/services/api.js. It has a class called ApiService with methods "
    "for every backend operation."
), size=11)

add_heading(doc, "How It Works", level=2)
api_flow = [
    "A page calls a function like api.analyzeKeyword({...})",
    "api.js checks if mock mode is on (VITE_USE_MOCK_DATA)",
    "If mock mode: returns fake data from built-in arrays (for testing)",
    "If real mode: sends an HTTP request to /api/n8n/webhook/[WorkflowName]",
    "Vercel's proxy forwards /api/n8n/... to https://vitcornu.app.n8n.cloud/...",
    "n8n processes the request and sends data back",
    "api.js transforms the response into the format the page expects",
    "The page gets the data and displays it"
]
for i, step in enumerate(api_flow, 1):
    add_bullet(doc, f"{i}. {step}")

add_heading(doc, "API Method Reference", level=2)
add_styled_table(doc,
    ["Method Name", "n8n Webhook", "HTTP Method", "What It Does"],
    [
        ["getBrands()", "/webhook/GetActiveBrands", "GET", "Fetches list of available brands"],
        ["getBrandDetails(name)", "/webhook/GetBrandDetails", "GET", "Gets full brand configuration"],
        ["analyzeKeyword(data)", "/webhook/AnalyzeKeyword_Backend", "POST", "Runs AI analysis, returns PAVE scores"],
        ["saveIdea(data)", "/webhook/SaveScoredIdea", "POST", "Saves scored keyword to SharePoint"],
        ["getBacklog(params)", "/webhook/FetchAllItems or FetchBriefedItem", "GET", "Gets content pipeline items (filtered by status)"],
        ["getContentDetails(id)", "/webhook/FetchItemDetails", "GET", "Gets full details for one item"],
        ["recommendFormat(data)", "/webhook/BriefRecommendation", "GET", "AI recommends a content format"],
        ["generateAuthorityBrief(data)", "/webhook/GenerateAuthorityBrief", "POST", "AI generates a full content brief"],
        ["updateBriefStatus(data)", "/webhook/UpdateBriefAndStatus", "PUT", "Saves approved brief to SharePoint"],
        ["produceContent(id)", "/webhook/ProduceContent", "POST", "Generates full HTML article"],
        ["publishContent(id)", "/webhook/UpdateContentStatus", "PUT", "Changes status to 'Published'"],
    ]
)

add_para(doc, "", size=6)
add_heading(doc, "The Mock Data System", level=2)
add_para(doc, (
    "api.js has built-in sample data (MOCK_BRANDS, MOCK_BACKLOG, etc.) so the "
    "website can be tested without needing the n8n backend running. "
    "When VITE_USE_MOCK_DATA is 'true', every API call returns this fake data "
    "with a realistic delay to simulate real behaviour."
), size=11)

add_note_box(doc, (
    "This is really handy for frontend dev. A developer can work on the UI "
    "without worrying about whether the backend's available."
), prefix="Tip: ")

# ===================================================================
#  SECTION 10: Adding a New Backend Connection
# ===================================================================

add_heading(doc, "10. Adding a New Backend Connection", level=1)

add_para(doc, (
    "If a new workflow gets added in n8n and you need the website to call it, "
    "here's what to do:"
), size=11)

add_heading(doc, "Step 1: Add the Method to api.js", level=2)
add_code_block(doc, (
    "// In src/services/api.js, add a new method to the ApiService class:\n\n"
    "async myNewFunction(data) {\n"
    "    if (USE_MOCK) {\n"
    "        await delay(1000)  // Simulate loading time\n"
    "        return { success: true, message: 'Mock response' }\n"
    "    }\n"
    "    // Real n8n call:\n"
    "    return this.request('/webhook/MyNewWorkflow', {\n"
    "        method: 'POST',\n"
    "        body: JSON.stringify({\n"
    "            field1: data.field1,\n"
    "            field2: data.field2\n"
    "        })\n"
    "    })\n"
    "}"
), label="Add to: src/services/api.js")

add_heading(doc, "Step 2: Call It from a Page", level=2)
add_code_block(doc, (
    "// In any page file (e.g., src/pages/MyPage.jsx):\n\n"
    "import { api } from '../services/api'\n\n"
    "const handleClick = async () => {\n"
    "    const result = await api.myNewFunction({ field1: 'value1' })\n"
    "    console.log(result)\n"
    "}"
), label="Use in: src/pages/[YourPage].jsx")

# ===================================================================
#  SECTION 11: Navigation System
# ===================================================================

add_heading(doc, "11. The Navigation System", level=1)

add_para(doc, "Navigation has two parts that need to stay in sync:", size=11)

add_styled_table(doc,
    ["Part", "File", "What It Controls"],
    [
        ["Routes", "src/App.jsx", "Which URL shows which page (the routing rules)"],
        ["Sidebar Links", "src/components/layout/Sidebar.jsx", "The clickable links in the left sidebar"],
    ]
)

add_para(doc, "", size=6)
add_para(doc, "Current navigation links:", size=11)
add_styled_table(doc,
    ["Name", "URL Path"],
    [
        ["Dashboard", "/dashboard"],
        ["Analyze Keyword", "/analyze"],
        ["Content Backlog", "/backlog"],
        ["Create Brief", "/brief"],
        ["Produce Content", "/produce"],
        ["Published", "/published"],
    ]
)

# ===================================================================
#  SECTION 12: Adding a New Navigation Link
# ===================================================================

add_heading(doc, "12. Adding a New Navigation Link", level=1)
add_para(doc, (
    "Same process as Section 7 (Adding a New Page), Step 3. "
    "Just add an entry to the navigation array in Sidebar.jsx."
), size=11)

add_code_block(doc, (
    "// Add to the navigation array in src/components/layout/Sidebar.jsx:\n"
    "{ name: 'Your New Page', href: '/your-url', icon: 'star' }"
), label="One-line change:")

# ===================================================================
#  SECTION 13: Reusable Components
# ===================================================================

add_heading(doc, "13. Reusable Components", level=1)

add_para(doc, (
    "These components are pre-built and can be dropped into any page. "
    "Just import and use."
), size=11)

add_heading(doc, "LoadingSpinner", level=2)
add_para(doc, "Shows a spinning animation with optional text. Use while data's loading.", size=11)
add_code_block(doc, (
    "import LoadingSpinner from '../components/common/LoadingSpinner'\n\n"
    "// Usage:\n"
    "<LoadingSpinner text=\"Loading data...\" />\n"
    "<LoadingSpinner size=\"lg\" text=\"Processing...\" />"
), label="How to use:")

add_heading(doc, "StatusBadge", level=2)
add_para(doc, "Shows a coloured pill with the item's status.", size=11)
add_code_block(doc, (
    "import StatusBadge from '../components/common/StatusBadge'\n\n"
    "// Usage:\n"
    "<StatusBadge status=\"Approved for Briefing\" />\n"
    "<StatusBadge status=\"Published\" />"
), label="How to use:")

add_heading(doc, "PAVEScoreCard", level=2)
add_para(doc, "The full PAVE score display with editable inputs and save button.", size=11)
add_code_block(doc, (
    "import PAVEScoreCard from '../components/common/PAVEScoreCard'\n\n"
    "// Usage:\n"
    "<PAVEScoreCard\n"
    "    scores={results.pave_scores}\n"
    "    searchSummary={results.search_summary}\n"
    "    onSave={handleSave}\n"
    "/>"
), label="How to use:")

add_heading(doc, "CSS Classes You Can Use Anywhere", level=2)
add_styled_table(doc,
    ["Class Name", "What It Does", "Usage Example"],
    [
        [".card", "A styled container box with border and padding", "<div className=\"card\">Content</div>"],
        [".btn btn-primary", "Purple filled button", "<button className=\"btn btn-primary\">Click</button>"],
        [".btn btn-secondary", "Outline/ghost button", "<button className=\"btn btn-secondary\">Click</button>"],
        [".btn-lg", "Makes a button larger", "Add to existing button classes"],
        [".btn-full", "Makes a button full-width", "Add to existing button classes"],
        [".animate-fade-in", "Fade-in animation on load", "<div className=\"animate-fade-in\">...</div>"],
        [".page-header", "Standard page title section", "<div className=\"page-header\"><h1>...</h1></div>"],
        [".page-title", "Large page heading style", "<h1 className=\"page-title\">Title</h1>"],
        [".page-subtitle", "Smaller subtitle under title", "<p className=\"page-subtitle\">Description</p>"],
        [".form-group", "Groups a label + input together", "<div className=\"form-group\">...</div>"],
        [".input", "Styled text input field", "<input className=\"input\" />"],
        [".select", "Styled dropdown select", "<select className=\"select\">...</select>"],
        [".label", "Styled form label", "<label className=\"label\">Name</label>"],
        [".badge badge-success", "Green status badge", "<span className=\"badge badge-success\">Done</span>"],
        [".badge badge-warning", "Amber status badge", "<span className=\"badge badge-warning\">Pending</span>"],
        [".table-container", "Scrollable table wrapper", "<div className=\"table-container\"><table>...</table></div>"],
        [".empty-state", "Centered empty message", "<div className=\"card empty-state\">No items</div>"],
        [".stats-grid", "Grid layout for stat cards", "<div className=\"stats-grid\">...</div>"],
        [".stat-card", "Individual statistic card", "<div className=\"stat-card\">...</div>"],
    ]
)

# ===================================================================
#  SECTION 14: Environment Settings
# ===================================================================

add_heading(doc, "14. Environment Settings", level=1)

add_styled_table(doc,
    ["Setting", "Purpose", "Current Value", "Where to Change"],
    [
        ["VITE_N8N_API_URL", "Where API calls go", "/api/n8n", ".env file + Vercel"],
        ["VITE_N8N_API_KEY", "n8n authentication key", "(stored securely)", ".env file + Vercel"],
        ["VITE_USE_MOCK_DATA", "Use fake data for testing?", "false", ".env file + Vercel"],
    ]
)

add_para(doc, "", size=6)
add_note_box(doc, (
    "Settings starting with VITE_ are special in Vite projects. They get baked into "
    "the website at build time. After changing a setting on Vercel, you'll need to "
    "redeploy for the change to take effect."
), prefix="How it works: ")

# ===================================================================
#  SECTION 15: Local Dev Setup (Detailed)
# ===================================================================

add_heading(doc, "15. Running the Website Locally", level=1)

add_para(doc, (
    "This section walks through everything a developer needs to get the website "
    "running on their own machine. We've included step-by-step instructions for "
    "Windows, Mac, and Linux since the commands are a bit different on each."
), size=11)

# --- 15.1 Prerequisites ---

add_heading(doc, "15.1 What You'll Need to Install", level=2)

add_para(doc, (
    "Before you can run the website locally, you need three things installed on your "
    "computer: Node.js (which includes npm), Git, and a code editor. Here's how to "
    "get each one."
), size=11)

# -- Node.js --

add_heading(doc, "Installing Node.js", level=3)

add_para(doc, (
    "Node.js is the runtime that lets you run JavaScript outside a browser. "
    "npm (Node Package Manager) comes bundled with it and is used to install "
    "all the project's dependencies. You want the LTS (Long Term Support) version, "
    "not the \"Current\" one."
), size=11)

add_para(doc, "On Windows:", size=11)
add_code_block(doc, (
    "1. Go to https://nodejs.org\n"
    "2. Click the big green LTS button to download the installer (.msi file)\n"
    "3. Run the installer, click Next through all the steps\n"
    "   - Leave all default options checked\n"
    "   - Make sure \"Add to PATH\" is checked (it should be by default)\n"
    "4. When it's done, open a NEW Command Prompt or PowerShell window\n"
    "5. Verify it worked by typing:"
), label="Windows:")
add_code_block(doc, (
    "node --version\n"
    "npm --version"
), label="Run these commands to verify (you should see version numbers):")

add_para(doc, "On Mac:", size=11)
add_code_block(doc, (
    "# Option A: Download from the website (easiest)\n"
    "1. Go to https://nodejs.org\n"
    "2. Click the LTS button to download the .pkg installer\n"
    "3. Run it, follow the prompts\n\n"
    "# Option B: Using Homebrew (if you have it installed)\n"
    "brew install node"
), label="Mac:")
add_code_block(doc, (
    "# Verify it worked:\n"
    "node --version\n"
    "npm --version"
))

add_para(doc, "On Linux (Ubuntu/Debian):", size=11)
add_code_block(doc, (
    "# Install Node.js using the NodeSource repository:\n"
    "curl -fsSL https://deb.nodesource.com/setup_lts.x | sudo -E bash -\n"
    "sudo apt-get install -y nodejs\n\n"
    "# Verify it worked:\n"
    "node --version\n"
    "npm --version"
), label="Linux (Ubuntu/Debian):")

add_note_box(doc, (
    "If you see version numbers like v20.x.x and 10.x.x, you're good. "
    "If you get \"command not found\", close and reopen your terminal, "
    "or restart your computer. The PATH sometimes doesn't update until you do."
), prefix="Troubleshooting: ")

# -- Git --

add_heading(doc, "Installing Git", level=3)

add_para(doc, (
    "Git is what we use to download the project code and push changes back "
    "to GitHub. It's the version control system."
), size=11)

add_para(doc, "On Windows:", size=11)
add_code_block(doc, (
    "1. Go to https://git-scm.com/download/win\n"
    "2. The download should start automatically (.exe file)\n"
    "3. Run the installer\n"
    "   - You can accept all default settings\n"
    "   - When it asks about the default editor, pick whatever you prefer\n"
    "     (VS Code if you have it, otherwise Notepad is fine)\n"
    "   - For \"Adjusting your PATH\", pick the recommended option\n"
    "     (\"Git from the command line and also from 3rd-party software\")\n"
    "4. Open a NEW Command Prompt or PowerShell window\n"
    "5. Verify:"
), label="Windows:")
add_code_block(doc, "git --version", label="Should show something like: git version 2.x.x")

add_para(doc, "On Mac:", size=11)
add_code_block(doc, (
    "# Git might already be installed. Check first:\n"
    "git --version\n\n"
    "# If it's not installed, Mac will prompt you to install\n"
    "# Xcode Command Line Tools. Click Install when prompted.\n\n"
    "# Or install via Homebrew:\n"
    "brew install git"
), label="Mac:")

add_para(doc, "On Linux (Ubuntu/Debian):", size=11)
add_code_block(doc, (
    "sudo apt-get update\n"
    "sudo apt-get install git\n\n"
    "# Verify:\n"
    "git --version"
), label="Linux:")

# -- Code Editor --

add_heading(doc, "Installing a Code Editor", level=3)

add_para(doc, (
    "You need a text editor to view and edit the code. VS Code is the most popular "
    "choice and works on all three operating systems. It's free."
), size=11)

add_styled_table(doc,
    ["Editor", "Download Link", "Notes"],
    [
        ["VS Code (Recommended)", "https://code.visualstudio.com", "Free, works on Windows/Mac/Linux, tons of extensions"],
        ["Cursor", "https://cursor.sh", "VS Code fork with built-in AI, great for AI-assisted coding"],
        ["Sublime Text", "https://www.sublimetext.com", "Lightweight and fast, free to evaluate"],
    ]
)

add_para(doc, "", size=6)
add_para(doc, (
    "Just download the installer for your OS from the website and run it. "
    "Nothing special to configure."
), size=11)

# --- 15.2 Opening a Terminal ---

add_heading(doc, "15.2 Opening a Terminal", level=2)

add_para(doc, (
    "You'll be typing commands into a terminal (also called command line or console). "
    "Here's how to open one on each OS:"
), size=11)

add_styled_table(doc,
    ["Operating System", "How to Open a Terminal"],
    [
        ["Windows", "Press Win + R, type 'cmd' and hit Enter. Or search for 'PowerShell' in the Start menu. Or in VS Code: press Ctrl + ` (backtick)"],
        ["Mac", "Open Finder > Applications > Utilities > Terminal. Or press Cmd + Space, type 'Terminal' and hit Enter. Or in VS Code: press Ctrl + ` (backtick)"],
        ["Linux", "Press Ctrl + Alt + T. Or in VS Code: press Ctrl + ` (backtick)"],
    ]
)

add_para(doc, "", size=6)
add_note_box(doc, (
    "If you're using VS Code or Cursor, the built-in terminal (Ctrl + ` backtick) "
    "is the easiest option. It opens right inside the editor so you don't have to "
    "switch windows."
), prefix="Tip: ")

# --- 15.3 Downloading and Running the Project ---

add_heading(doc, "15.3 Downloading and Running the Project", level=2)

add_para(doc, (
    "Now that you've got Node.js, Git, and a code editor installed, here's how to "
    "get the project running. These commands are the same on Windows, Mac, and Linux."
), size=11)

add_para(doc, "Step 1: Clone (download) the code from GitHub", size=11)
add_code_block(doc, (
    "git clone https://github.com/ifrahimafzal808-sudo/AI_Content.git"
), label="Run this in your terminal:")

add_para(doc, (
    "This creates a folder called AI_Content with all the project files inside it."
), size=11)

add_para(doc, "Step 2: Navigate into the frontend folder", size=11)
add_code_block(doc, (
    "cd AI_Content/frontend"
), label="On Windows, Mac, and Linux:")

add_note_box(doc, (
    "On Windows, you can also use backslashes: cd AI_Content\\frontend. "
    "Both work."
), prefix="Windows note: ")

add_para(doc, "Step 3: Install all the project dependencies", size=11)
add_code_block(doc, (
    "npm install"
), label="This downloads everything the project needs:")

add_para(doc, (
    "This reads the package.json file and downloads all the required libraries "
    "into a folder called node_modules. It might take a minute or two the first "
    "time. You'll see a progress bar. Don't worry about any \"warn\" messages, "
    "those are usually fine."
), size=11)

add_para(doc, "Step 4: Set up your environment file", size=11)
add_para(doc, (
    "The project needs a .env file with some settings. There's a template file "
    "called .env.example that you can copy:"
), size=11)

add_code_block(doc, (
    "# On Windows (Command Prompt):\n"
    "copy .env.example .env\n\n"
    "# On Windows (PowerShell):\n"
    "Copy-Item .env.example .env\n\n"
    "# On Mac/Linux:\n"
    "cp .env.example .env"
), label="Copy the template:")

add_para(doc, (
    "Then open the .env file in your code editor and fill in the values. "
    "The most important setting for local dev is:"
), size=11)
add_code_block(doc, (
    "# Set this to 'true' to use fake test data (no backend needed)\n"
    "VITE_USE_MOCK_DATA=true\n\n"
    "# Or set to 'false' to use the real n8n backend\n"
    "VITE_USE_MOCK_DATA=false\n"
    "VITE_N8N_API_URL=/api/n8n\n"
    "VITE_N8N_API_KEY=your_api_key_here"
), label="Key settings in .env:")

add_para(doc, "Step 5: Start the development server", size=11)
add_code_block(doc, (
    "npm run dev"
), label="Run this command:")

add_para(doc, (
    "You'll see output like this in the terminal:"
), size=11)
add_code_block(doc, (
    "  VITE v7.x.x  ready in xxx ms\n\n"
    "  >  Local:   http://localhost:5173/\n"
    "  >  Network: http://192.168.x.x:5173/"
))

add_para(doc, "Step 6: Open the website in your browser", size=11)
add_para(doc, (
    "Open your web browser (Chrome, Firefox, Edge, Safari, whatever you use) and go to:"
), size=11)
add_para(doc, "http://localhost:5173", bold=True, size=12, color=(79, 70, 229))

add_para(doc, (
    "You should see the AI Content Copilot landing page. That's it, you're running "
    "the website locally!"
), size=11)

# --- 15.4 Dev Workflow Tips ---

add_heading(doc, "15.4 Day-to-Day Development Tips", level=2)

add_para(doc, "Things that'll make your life easier while developing:", size=11)

add_styled_table(doc,
    ["What", "How"],
    [
        ["See changes instantly", "The dev server watches your files. Save a file and the browser refreshes automatically. No need to restart anything"],
        ["Test without the backend", "Set VITE_USE_MOCK_DATA=true in your .env file. The website will use built-in fake data so you can work on the UI without needing n8n running"],
        ["Stop the dev server", "Go to your terminal and press Ctrl+C (works on all operating systems)"],
        ["Restart the dev server", "Press Ctrl+C to stop it, then run 'npm run dev' again"],
        ["Install a new library", "Run 'npm install library-name' (e.g., npm install axios). This updates package.json automatically"],
        ["Build for production", "Run 'npm run build' to create the optimized files that go to Vercel. The output goes into a 'dist' folder"],
        ["Check for errors", "Open your browser's developer tools (F12 on most browsers > Console tab) to see JavaScript errors"],
    ]
)

add_para(doc, "", size=6)

# --- 15.5 Common Problems ---

add_heading(doc, "15.5 Common Problems When Setting Up Locally", level=2)

add_styled_table(doc,
    ["Problem", "Cause", "Fix"],
    [
        ["'node' is not recognized / command not found", "Node.js isn't installed or not in your PATH", "Reinstall Node.js, make sure 'Add to PATH' is checked. Close and reopen your terminal after installing"],
        ["'git' is not recognized / command not found", "Git isn't installed or not in your PATH", "Reinstall Git, accept the default PATH option. Close and reopen your terminal"],
        ["npm install shows errors", "Usually network issues or permission problems", "Try running as administrator (Windows) or with sudo (Mac/Linux). Or try 'npm cache clean --force' first"],
        ["Port 5173 is already in use", "Another dev server or app is using that port", "Either close the other app, or run 'npm run dev -- --port 3000' to use a different port"],
        ["Blank page after npm run dev", "Usually a JavaScript error or wrong .env config", "Check the terminal for errors. Check the browser console (F12). Make sure .env file exists"],
        ["API calls fail locally", "Mock mode is off but backend isn't accessible", "Set VITE_USE_MOCK_DATA=true in .env for local testing, or make sure n8n Cloud is running"],
    ]
)

# ===================================================================
#  SECTION 16: Deployment
# ===================================================================

add_heading(doc, "16. How Deployments Work", level=1)

deploy_flow = [
    "Developer makes changes locally",
    "Developer pushes to GitHub ('git push')",
    "Vercel picks up the push automatically",
    "Vercel runs 'npm run build' to compile",
    "If successful, the new version goes live",
    "If it fails, the old version stays up (no downtime)",
]
for i, step in enumerate(deploy_flow, 1):
    add_bullet(doc, f"{i}. {step}")

add_para(doc, "", size=6)
add_note_box(doc, (
    "Vercel never takes down the old version until the new one's confirmed working. "
    "If a deployment fails, your website just keeps running the previous version. "
    "There's zero risk of downtime from a bad deploy."
), prefix="Safety net: ")

# ===================================================================
#  SECTION 17: Maintaining with AI
# ===================================================================

add_heading(doc, "17. Maintaining the App with AI", level=1)

add_para(doc, (
    "Because the project's using standard React with clean, well-organized code, "
    "AI coding assistants do really well with it."
), size=11)

add_heading(doc, "AI Coding Tools Worth Looking At", level=2)
add_styled_table(doc,
    ["Tool", "Best For", "Cost"],
    [
        ["Claude (Anthropic)", "Complex code changes, understanding architecture", "Free tier available, Pro ~$20/month"],
        ["GPT-5 (OpenAI / ChatGPT)", "General code changes, explanations", "Free tier available, Plus ~$20/month"],
        ["Cursor (IDE)", "Full dev environment with built-in AI", "Free tier, Pro ~$20/month"],
        ["GitHub Copilot", "Inline code suggestions while typing", "~$10-19/month"],
        ["Lovable", "No-code frontend generation through prompting", "Subscription-based, good for simpler UIs"],
    ]
)

add_para(doc, "", size=6)
add_heading(doc, "How to Prompt an AI Assistant", level=2)
add_para(doc, "When asking an AI to make changes, include these details:", size=11)

prompt_tips = [
    ("Tell it the stack:", "\"This is a React + Vite project using vanilla CSS\""),
    ("Point to the file:", "\"Edit src/pages/AnalyzeKeyword.jsx\""),
    ("Say what you want:", "\"Add a text field called 'Notes' below the keyword input\""),
    ("Mention styling:", "\"Use the existing CSS classes: className='form-group' for the wrapper and className='input' for the field\""),
    ("Give context:", "\"The API service is in src/services/api.js and all styles are in src/index.css\""),
]
for prefix, example in prompt_tips:
    p = doc.add_paragraph()
    run_b = p.add_run(prefix + " ")
    run_b.font.size = Pt(10)
    run_e = p.add_run(example)
    run_e.italic = True
    run_e.font.size = Pt(10)
    run_e.font.color.rgb = RGBColor(107, 114, 128)

add_para(doc, "", size=4)
add_code_block(doc, (
    "EXAMPLE PROMPT:\n\n"
    "\"I've got a React + Vite frontend for an AI content platform.\n"
    "Here's the structure:\n"
    "  - Pages are in src/pages/\n"
    "  - API calls are in src/services/api.js\n"
    "  - Styles are in src/index.css\n"
    "  - Routes are defined in src/App.jsx\n"
    "  - Sidebar links are in src/components/layout/Sidebar.jsx\n\n"
    "I want to add a new page called 'Analytics' at /analytics that\n"
    "shows a chart of how many articles were published per week.\n"
    "Please create the page file, add the route, and add the sidebar link.\""
), label="Example Prompt:")

# ===================================================================
#  SECTION 18: Quick Reference Cheat Sheet
# ===================================================================

add_heading(doc, "18. Quick Reference Cheat Sheet", level=1)

add_para(doc, "Keep this page handy:", size=11)

add_styled_table(doc,
    ["I Want To...", "Edit This File", "Section"],
    [
        ["Change the website colours", "src/index.css (CSS variables at top)", "Section 5"],
        ["Change a page's title or text", "src/pages/[PageName].jsx", "Section 8"],
        ["Add a brand-new page", "New file in src/pages/ + App.jsx + Sidebar.jsx", "Section 7"],
        ["Add a sidebar link", "src/components/layout/Sidebar.jsx", "Section 12"],
        ["Connect a new n8n workflow", "src/services/api.js", "Section 10"],
        ["Change the navbar logo/title", "src/components/layout/Navbar.jsx", "Section 8"],
        ["Toggle mock data on/off", ".env file > VITE_USE_MOCK_DATA", "Section 14"],
        ["Change the n8n backend URL", ".env + vercel.json + vite.config.js", "Section 14"],
        ["Deploy a new version", "Push code to GitHub (automatic)", "Section 16"],
        ["Force a redeployment", "Vercel dashboard > Deployments > Redeploy", "Section 16"],
    ]
)

# Footer
doc.add_paragraph()
divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run("_" * 50)
run.font.color.rgb = RGBColor(209, 213, 219)


# ===================================================================
#  SAVE
# ===================================================================

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "AI_Content_Copilot_Frontend_Maintenance_Guide.docx"
)
doc.save(output_path)
print(f"\n  Document saved to:\n    {output_path}")
print(f"    File size: {os.path.getsize(output_path) / 1024:.1f} KB")
